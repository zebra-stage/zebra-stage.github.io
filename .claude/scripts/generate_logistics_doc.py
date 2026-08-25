"""Generate the Sri Lanka shipment logistics process document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(r"G:\OneDrive - Zebra Technologies\01 Administration\06 Company Policies & Guidelines\Logistics")
OUT_FILE = OUT_DIR / "Sri Lanka Shipment - End-to-End Process.docx"

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Helpers ---------------------------------------------------------------------
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def add_bullets(items, level=0):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.63 + 0.63 * level)

def add_numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)

def add_table(headers, rows, header_fill="1F3A5F", widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    if widths:
        for col_idx, w in enumerate(widths):
            for row in t.rows:
                row.cells[col_idx].width = w
    return t

def add_reference_callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("Reference: ")
    r.bold = True
    r.italic = True
    r2 = p.add_run(text)
    r2.italic = True

# Title -----------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Sri Lanka Shipment — End-to-End Process")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Ordering, Commercial Invoice Approval, Tracking and Hand-off to Customs Clearance")
sr.italic = True
sr.font.size = Pt(12)

doc.add_paragraph()

# Metadata table --------------------------------------------------------------
meta = doc.add_table(rows=3, cols=2)
meta.style = "Light List Accent 1"
meta_rows = [
    ("Document Owner", "Eranga Tennakoon — eranga.tennakoon@zebra.com"),
    ("Audience", "Zebra Sri Lanka engineers and project owners initiating international shipments"),
    ("Status", "Draft"),
]
for i, (k, v) in enumerate(meta_rows):
    c0 = meta.rows[i].cells[0]
    c0.text = ""
    rr = c0.paragraphs[0].add_run(k)
    rr.bold = True
    meta.rows[i].cells[1].text = v

doc.add_paragraph()

# 1. Purpose ------------------------------------------------------------------
add_heading("1. Purpose", level=1)
add_para(
    "This document describes the end-to-end process for receiving devices and "
    "equipment at the Zebra Sri Lanka office through international shipment, "
    "and for exporting devices out of Sri Lanka. It covers everything an "
    "engineer or project owner must do — from identifying the need, raising "
    "the commercial invoice with the counterpart, obtaining pre-shipment "
    "approvals, despatch and tracking, through TRC/ICL licensing and customs "
    "clearance in Sri Lanka, to delivery at site."
)
add_para(
    "TRC submission mechanics and Zebra's master routing/shipping guidelines "
    "are not duplicated here — they are owned by the reference documents in "
    "Section 14. This document is otherwise self-contained."
)

# 2. Scope --------------------------------------------------------------------
add_heading("2. Scope", level=1)
add_bullets([
    "Inbound shipments to the Zebra Sri Lanka office (Colombo) initiated by engineers or project owners.",
    "Outbound (export) shipments from Sri Lanka.",
    "Both Free-of-Charge (FOC) intercompany transfers and supplier-sold shipments.",
    "Devices that may contain radios (WWAN, WLAN, BT) and therefore require TRC clearance.",
    "Hand-off from the requester to the in-country Logistics Analyst (Harsha Mendis) and onward to the broker, TRC, ICL and Sri Lanka Customs.",
])
add_para("Out of scope:", bold=True)
add_bullets([
    "End-of-life device disposal and e-waste recycling — covered by a separate process.",
    "Day-to-day broker operational mechanics not visible to the requester.",
])

# 3. Roles & Contacts ---------------------------------------------------------
add_heading("3. Roles and Contacts", level=1)
add_para(
    "The table below identifies the people and groups involved end-to-end. "
    "All commercial invoice approvals must include the pre-approval distribution list."
)
add_table(
    headers=["Role", "Responsibility", "Contact"],
    rows=[
        ["Requester / Engineer", "Identifies the device need, owns the cost centre, initiates the request with the counterpart, and is notified of tracking once shipped.", "(this document's author / project owner)"],
        ["Counterpart (Shipper)", "Raises the Commercial Invoice, prepares Packing List, Bill of Lading / Air Waybill and Certificate of Origin where applicable.", "Originating Zebra entity or external supplier"],
        ["Logistics Analyst — Sri Lanka", "Reviews and approves the Commercial Invoice (HS code, item description, value, shipping terms), obtains Import Control approval where required, owns in-country clearance hand-off.", "Harsha Mendis — harsha.mendis@zebra.com"],
        ["Logistics Analyst — Sri Lanka (Backup)", "Pre-approval distribution / backup reviewer for commercial invoices.", "Yashodha Hansini Godagama — YHansini@zebra.com"],
        ["APAC Logistics", "Regional logistics review and pre-approval.", "GTCO-APAC@zebra.com; APAC-Logistics@zebra.com"],
        ["Customs Broker", "Collects shipment documents, prepares cargo held notification, submits to TRC / ICL / SL Customs.", "Assigned by Logistics Analyst"],
        ["TRC (Telecommunications Regulatory Commission)", "Issues TRC licence for radio-bearing devices.", "Via broker"],
        ["ICL (Import Control Licence)", "Issues import licence for restricted goods.", "Via broker"],
        ["Sri Lanka Customs", "Clears the goods on receipt of declaration and licences.", "Via broker"],
    ],
)

# 4. Process Overview ---------------------------------------------------------
add_heading("4. End-to-End Process Overview", level=1)
add_para(
    "The full inbound lifecycle has seven stages. Stages 1–4 are owned "
    "primarily by the requesting engineer and the Logistics Analyst. "
    "Stages 5–7 are executed by the broker, TRC/ICL and Sri Lanka Customs. "
    "Outbound (export) shipments follow a shorter four-step procedure (§ 13)."
)
add_table(
    headers=["#", "Stage", "Primary Owner", "Section"],
    rows=[
        ["1", "Identify need and engage counterpart", "Requester", "§ 6"],
        ["2", "Commercial Invoice prepared by counterpart", "Counterpart / Shipper", "§ 7"],
        ["3", "Commercial Invoice review and pre-shipment approvals", "Logistics Analyst (Harsha)", "§ 8"],
        ["4", "Shipment dispatch and tracking notification", "Counterpart → Requester + Logistics", "§ 9"],
        ["5", "TRC / ICL licensing (if required)", "Broker / TRC / ICL", "§ 11"],
        ["6", "Customs clearance in Sri Lanka", "Broker / SL Customs", "§ 12"],
        ["7", "Delivery to Zebra Sri Lanka office", "Deliverer", "§ 12"],
    ],
)

# 5. Key Principles -----------------------------------------------------------
add_heading("5. Key Principles for Smooth Customs Clearance", level=1)
add_para(
    "Three principles underpin every shipment and should be checked early — "
    "before a Commercial Invoice is raised, not after the goods arrive."
)
add_numbered([
    "Licensing: Proactively identify any products that require special import/export licences or permits from Sri Lankan government agencies (TRC, ICL, BOI, Customs). Apply for these well in advance of shipment.",
    "Tariff Codes (HS Codes): Ensure the correct Harmonized System (HS) code is assigned to every product. This is critical for determining the correct import duties and taxes.",
    "Accurate Documentation: Prepare all required trade documents with precise and complete information in accordance with Sri Lankan legislation.",
])

# 5. Stage 1: Identify need ---------------------------------------------------
add_heading("6. Stage 1 — Identify Need and Engage Counterpart", level=1)
add_para("Trigger:", bold=True)
add_para(
    "Engineer or project owner identifies a need for devices, accessories or "
    "equipment that must be shipped to the Sri Lanka office from another Zebra "
    "entity or an external supplier."
)
add_para("Actions:", bold=True)
add_numbered([
    "Confirm the list of items required: part numbers, models, quantities, intended purpose (e.g. testing, development, demo) and the cost centre to be billed.",
    "Identify the counterpart who will ship the items (Zebra entity owner or supplier contact).",
    "Loop the Logistics Analyst (Harsha Mendis) into the initial discussion early. This is critical — it allows Harsha to flag any licensing, HS code, or Import Control requirements before a Commercial Invoice is raised.",
    "Share the device specification, including radio frequency bands and transmission power levels for any product containing or considered a radio (WWAN, WLAN, BT). This is required so that Sri Lanka can confirm the products are certifiable for import.",
    "Request the counterpart to prepare a draft Commercial Invoice (see § 6).",
])
add_para("Key checks before requesting an invoice:", bold=True)
add_bullets([
    "For radio-transmitting devices like mobile computers, a single shipment cannot contain more than 3 devices. If more than 3 units are required, plan to split into separate shipments of ≤3 units each.",
    "For Free-of-Charge (FOC) shipments, Import Control approval must be obtained BEFORE shipment arrival. The shipper must wait for the green light from Zebra Sri Lanka before despatching.",
    "Shipments valued over USD 1,000 require Zebra Sri Lanka to apply for import approval with the Department of Import and Export Control before the shipment departs origin.",
])
add_reference_callout(
    "Shipping to Sri Lanka 2018 Sep 12th.pdf — \"Prior to Shipping\" section, for the device specification, FOC pre-approval, and USD 1,000 import approval rules."
)

# 6. Stage 2: Commercial Invoice ---------------------------------------------
add_heading("7. Stage 2 — Commercial Invoice Prepared by Counterpart", level=1)
add_para(
    "The counterpart prepares a draft Commercial Invoice. The invoice must "
    "carry the correct \"Ship To\", \"Bill To\", FOC declaration statement, "
    "HS codes, and shipping terms. The exact text and address blocks are "
    "specified in the Shipping to Sri Lanka guideline."
)
add_para("Required content on the Commercial Invoice:", bold=True)
add_bullets([
    "Ship To: Zebra Technologies Lanka (Pvt) Ltd, Maga Towers, Colombo 05 — full address per Shipping to Sri Lanka PDF.",
    "Bill To: depends on whether the shipment is supplier-sold to Zebra, Free-of-Charge, or from Zebra Singapore — see PDF for the three variants.",
    "For FOC shipments, the prescribed declaration statement: \"These items are fully paid by Zebra Technologies Corporation, USA & shipped to its Sri Lankan subsidiary, Zebra Technologies Lanka (Pvt) Ltd entirely free of charge. Not for resale. No commercial value. Intercompany use of testing and development only.\"",
    "Accurate HS (Harmonized System) code per item.",
    "Item description, quantity, unit value and total value (must reflect authentic values).",
    "Seller, buyer, currency, country of origin.",
    "Shipping terms / Incoterms per the table in the Shipping to Sri Lanka guideline.",
])
add_reference_callout(
    "Shipping to Sri Lanka 2018 Sep 12th.pdf — \"Commercial Invoice\" section for exact address blocks, Bill To variants, and the FOC declaration wording. "
    "Zebra Routing Guide 10.5.pdf — for Incoterms, FedEx account numbers, and inbound-to-APAC carrier instructions."
)

# 7. Stage 3: Approval -------------------------------------------------------
add_heading("8. Stage 3 — Commercial Invoice Review and Pre-Shipment Approvals", level=1)
add_para(
    "A draft copy of the Commercial Invoice must be emailed to the pre-approval "
    "distribution list BEFORE the shipment is despatched. The Logistics Analyst "
    "reviews and may request changes; the final approved copy is what must "
    "accompany the physical shipment."
)
add_para("Pre-approval distribution list (must be on the email):", bold=True)
add_bullets([
    "Kumara Mendis, Harsha — harsha.mendis@zebra.com",
    "Yashodha Hansini Godagama — YHansini@zebra.com",
    "GTCO-APAC@zebra.com",
    "APAC-Logistics@zebra.com",
])
add_para("What the Logistics Analyst verifies:", bold=True)
add_table(
    headers=["Item", "What is checked", "Why it matters"],
    rows=[
        ["HS Code", "Each item's Harmonized System code matches the device.", "Determines import duty, tax band and any restricted-goods licensing."],
        ["Item description", "Description is unambiguous and matches the physical goods.", "Mismatched descriptions cause customs to hold the cargo."],
        ["Value", "Stated unit and total value is authentic.", "Customs reject undervaluation; valuation drives duty."],
        ["Shipping terms / Incoterms", "Incoterm matches the Bill To / Free-Of-Charge arrangement.", "Wrong Incoterm shifts duty/tax responsibility unexpectedly."],
        ["Ship To / Bill To blocks", "Match Zebra Sri Lanka / Zebra Lincolnshire / Singapore template.", "Required for the consignee to clear the shipment."],
        ["FOC declaration", "Present and verbatim for FOC shipments.", "Customs needs the statement to process FOC without commercial value."],
        ["Radio device count", "≤3 radio devices per shipment.", "More than 3 triggers a TRC restriction and shipment delay."],
        ["Import Control pre-approval", "Required when value > USD 1,000 or for FOC shipments.", "Without it, shipments will be held or penalised on arrival."],
    ],
)
add_para("Outcome:", bold=True)
add_bullets([
    "Logistics Analyst returns the final copy of the Commercial Invoice to the counterpart.",
    "Counterpart must use this final approved copy in the shipment.",
    "Green light to ship is issued by Zebra Sri Lanka.",
])
add_reference_callout(
    "Shipping to Sri Lanka 2018 Sep 12th.pdf — pre-approval recipient list and green-light requirement."
)

# 8. Stage 4: Dispatch & Tracking --------------------------------------------
add_heading("9. Stage 4 — Shipment Dispatch and Tracking Notification", level=1)
add_para("Actions:", bold=True)
add_numbered([
    "Counterpart despatches the shipment using the approved Commercial Invoice and the Incoterms agreed in the routing guide.",
    "Immediately after despatch, the counterpart sends the tracking number (e.g. FedEx AWB) to BOTH the Logistics Analyst pre-approval distribution list AND the original requester / engineer who initiated the order.",
    "The Logistics Analyst registers the inbound shipment and engages the broker to prepare cargo-held notification documents in advance of arrival.",
    "Requester monitors the tracking and confirms expected arrival date with the Logistics Analyst.",
])
add_para("Notification email — minimum content:", bold=True)
add_bullets([
    "Carrier and tracking number.",
    "Despatch date and expected arrival date.",
    "Number of packages, gross weight and dimensions.",
    "Reference to the approved Commercial Invoice (attach final PDF).",
    "Cost centre / billing reference, if FOC.",
])
add_reference_callout(
    "Zebra Routing Guide 10.5.pdf — carrier accounts, lanes and \"Inbound to APAC\" rules; "
    "Shipping to Sri Lanka 2018 Sep 12th.pdf — \"After ship out, sender provide the tracking# to above pre-approval list.\""
)

# 10. Required Documents for Import ------------------------------------------
add_heading("10. Required Documents for Import", level=1)
add_para(
    "The following documents must be prepared for the Sri Lanka Customs "
    "declaration. The shipper is responsible for the trade documents; the "
    "Logistics Analyst is responsible for in-country licences and permits."
)
add_table(
    headers=["Document", "Notes", "Responsible Party"],
    rows=[
        ["Bill of Lading / Air Waybill", "Bill of Lading for sea freight; Air Waybill for air freight.", "Shipper"],
        ["Commercial Invoice", "Must show authentic values, seller, buyer, currency, country of origin. Final approved copy only (see § 8).", "Shipper"],
        ["Packing List", "Itemised contents matching the Commercial Invoice.", "Shipper"],
        ["Certificate of Origin", "Required when claiming preferential origin or when destination customs request it.", "Shipper / Manufacturer"],
        ["Import / Export / TRC / BOI / Customs licences or permits", "Restricted-goods licences (e.g. TRC for radio devices, ICL for Import Control).", "Logistics Analyst — Harsha"],
        ["Other certificates / supporting documents", "Product specification, item details and stated purpose of import, and any other documents requested by Sri Lanka Customs.", "Requester / Shipper"],
    ],
)

# 11. Stage 5: TRC / ICL Licensing -------------------------------------------
add_heading("11. Stage 5 — TRC and ICL Licensing", level=1)
add_para(
    "Where the shipment contains radio-bearing devices or restricted goods, "
    "the broker submits documents and payment to the relevant authority to "
    "obtain the licence before the goods can be released by Sri Lanka Customs."
)
add_para("TRC submission — main documents required:", bold=True)
add_bullets([
    "TRC Request Letter",
    "TRC Application",
    "Annexure One",
    "Invoice Copy (with TRC seal on the final approved copy)",
    "Device Specification",
])
add_para("Licensing timelines:", bold=True)
add_table(
    headers=["Licence", "Responsible", "Timeline"],
    rows=[
        ["TRC Licence (radio devices)", "Broker / TRC", "4 working days"],
        ["ICL Licence (Import Control)", "Broker / ICL", "1 working day"],
    ],
)
add_reference_callout(
    "TRC Submission and Clearance Procedure.pdf — full TRC submission checklist, payment, FedEx authorisation letter, cargo held notification form, and clearance flow."
)

# 12. Stage 6 & 7: Customs Clearance and Delivery ----------------------------
add_heading("12. Stages 6 & 7 — Customs Declaration, Clearance and Delivery", level=1)
add_para(
    "Two parallel views of the same activity follow. The first table is the "
    "Zebra-specific workflow as executed by our broker and Logistics Analyst. "
    "The second table is the Sri Lanka Customs procedural view — what happens "
    "behind the scenes at the port or airport."
)

add_heading("12.1 Zebra-specific clearance workflow (imports)", level=2)
add_table(
    headers=["#", "Action", "Responsible", "Notes / Timeline"],
    rows=[
        ["1", "Package received and held on arrival", "SL Customs", "The process begins when the shipment arrives in country."],
        ["2", "Broker collects shipment documents and delivers them to the Zebra office", "Broker", "The broker acts as the intermediary."],
        ["3", "Prepare request letter to customs and cargo held notification form; submit to broker", "Zebra office (Logistics Analyst)", "Formalises the clearance request."],
        ["4", "TRC Licence: submit documents and payment to TRC", "Broker / TRC", "4 working days."],
        ["5", "ICL Licence: submit documents and payment to ICL", "Broker / ICL", "1 working day."],
        ["6", "Submit all documents (incl. licences) to SL Customs and complete package clearance", "Broker / SL Customs", "This step formally clears the goods through customs."],
        ["7", "Deliver cleared package to the Zebra Sri Lanka site", "Deliverer", "The package is transported to the Zebra site."],
    ],
)

add_heading("12.2 Sri Lanka Customs procedural steps (Direct Import Procedure)", level=2)
add_numbered([
    "Submit Manifest: the shipping carrier or consolidator submits the cargo manifest to Sri Lanka Customs through the designated electronic system.",
    "De-grouping (if applicable): the carrier agent separates the master Bill of Lading into individual shipments.",
    "Lodge Customs Declaration: a Customs Broker or authorised agent completes and submits the official cargo declaration (CUSDEC) along with all trade documents to Sri Lanka Customs via their Electronic Data Interchange (EDI) system.",
    "Pay Duties and Fees: after the declaration is assessed by Customs, a payment voucher is generated. Pay all applicable customs duties, taxes and fees.",
    "Customs Clearance and Inspection: facilitate the customs clearance process at the port or airport. This may include a physical examination of the goods to verify their nature, quantity and value against the declaration.",
    "Goods Released: once all port and customs formalities are complete and duties are paid, the goods are cleared and can be transported to the warehouse.",
    "Post-Clearance Audit: Sri Lanka Customs may conduct a post-clearance audit at a later date to verify compliance. Retain all shipment documentation accordingly.",
])

# 13. Export Procedure -------------------------------------------------------
add_heading("13. Export Procedure (Outbound from Sri Lanka)", level=1)
add_para(
    "Outbound shipments from Sri Lanka follow a shorter procedure. The "
    "requester engages the Logistics Analyst as early as for imports so that "
    "export licences (where applicable) can be obtained before the goods are "
    "ready to ship."
)
add_para("Required documents for export:", bold=True)
add_bullets([
    "Bill of Lading or Air Waybill",
    "Pro-Forma Invoice or Commercial Invoice",
    "Customs Packing List",
    "Export Licence or Permit (if applicable)",
    "Any required Health or Quality Certificates",
])
add_para("Step-by-step export process:", bold=True)
add_numbered([
    "Prepare Documents: prepare all necessary export documents and obtain any required export licences from the relevant Sri Lankan government departments.",
    "Submit Export Declaration: submit the official export declaration to Sri Lanka Customs through the designated electronic system.",
    "Customs Clearance: facilitate the customs export clearance at the port. This may include a customs inspection.",
    "Goods Cleared for Export: once all formalities are complete, the goods are cleared for export and can be loaded onto the vessel or aircraft.",
])

# 14. Reference Documents ----------------------------------------------------
add_heading("14. Reference Documents", level=1)
add_para(
    "All documents referenced below live in: "
    "G:\\OneDrive – Zebra Technologies\\01 Administration\\06 Company Policies & Guidelines\\Logistics\\"
)
add_table(
    headers=["Document", "Purpose / What to use it for"],
    rows=[
        ["Shipping to Sri Lanka 2018 Sep 12th.pdf", "Pre-shipment rules: device specification approval, FOC Import Control pre-approval, USD 1,000 rule, Commercial Invoice address blocks, FOC declaration statement, pre-approval distribution list, radio-device shipment limit, Incoterms guidance."],
        ["TRC Submission and Clearance Procedure.pdf", "Step-by-step TRC submission: required documents, payment, TRC Approval Letter handling, FedEx authorisation letter, cargo held notification form, and the imports clearance flow diagram."],
        ["Zebra Routing Guide 10.5.pdf", "Master routing/shipping guidelines for suppliers — inbound to APAC lanes, FedEx account numbers, ASN labels, IATA marks, dangerous goods declarations, ship method instructions."],
    ],
)

# 15. Open Items -------------------------------------------------------------
add_heading("15. Open Items", level=1)
add_para(
    "Items to be confirmed or added as this document is socialised:"
)
add_bullets([
    "Sample documents (Commercial Invoice, Packing List, Certificate of Origin, TRC Request Letter, Cargo Held Notification Form) to be linked once approved exemplars are available.",
    "Standard email templates for: (a) initial counterpart request, (b) Commercial Invoice pre-approval submission, (c) tracking notification to Logistics + requester.",
    "Confirm the current broker name and contact (referenced as \"Broker\" in the SOP).",
    "Confirm whether ICL / Import Control rules from the 2018 Shipping to Sri Lanka guideline are still current, or refresh from the Department of Import and Export Control.",
])

# 16. Revision History -------------------------------------------------------
add_heading("16. Revision History", level=1)
add_table(
    headers=["Version", "Date", "Author", "Change"],
    rows=[
        ["0.1", "2026-05-18", "Eranga Tennakoon", "Initial draft consolidating ordering workflow, Commercial Invoice approval, tracking notification, and references to existing SOP / TRC / Routing / Shipping guideline / E-waste documents."],
        ["0.2", "2026-05-18", "Eranga Tennakoon", "Absorbed content of Shipment Standard Operating Procedure.docx (Key Principles, Required Documents for Import, Direct Import seven-step process, Export Procedure). Removed SOP and E-waste from reference list. Renumbered sections."],
    ],
)

doc.save(OUT_FILE)
print(f"Wrote: {OUT_FILE}")
